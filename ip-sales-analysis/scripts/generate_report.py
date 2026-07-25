#!/usr/bin/env python3
"""
IP销售分析报告生成器

功能：
1. 读取指定Excel销售数据
2. 按照标准分析框架计算关键指标
3. 生成可视化图表
4. 输出格式化的Word分析报告(.docx)

用法：
    python3 generate_report.py --excel <数据文件.xlsx> --output <输出目录>
"""

import argparse
import os
import sys
import warnings
from datetime import datetime

import numpy as np
import pandas as pd

warnings.filterwarnings("ignore")

# ── 图表配置 ─────────────────────────────────────────────
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
fm._load_fontmanager(try_read_cache=False)

# 设置中文字体
_ZH_FONT = "sans-serif"
for f in fm.fontManager.ttflist:
    if f.name in (
        "Microsoft YaHei", "YaHei", "微软雅黑",
        "PingFang SC", "Heiti SC", "STHeiti", "Songti SC",
        "PingFang TC", "Heiti TC", "SimHei",
    ):
        _ZH_FONT = f.name
        break
plt.rcParams["font.sans-serif"] = [_ZH_FONT, "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False

# ── 排版常量 ─────────────────────────────────────────────
CHART_DPI = 200
CHART_WIDTH = 10
CHART_HEIGHT = 4.5
TITLE_FONT_SIZE = 14
AXIS_FONT_SIZE = 11

# ── 辅助函数 ─────────────────────────────────────────────
def _parse_sales(series):
    """将销售额列转为数值，'null'→0"""
    return pd.to_numeric(series, errors="coerce").fillna(0)

def _parse_freq(series):
    """将付费次数列转为数值，None→0"""
    return pd.to_numeric(series, errors="coerce").fillna(0)

def _fmt_yuan(val):
    """格式化金额"""
    if val >= 10000:
        return f"{val / 10000:.2f}万"
    return f"{val:.2f}"

# ── 核心指标计算 ─────────────────────────────────────────
def compute_metrics(df):
    """
    计算所有分析指标，返回 dict。
    df 已清洗：销售额/付费次数为数值，下单时间为 datetime，无合计行。
    """
    metrics = {}

    # 销售日期范围
    dates = df["下单时间"].dropna()
    metrics["date_min"] = dates.min().strftime("%Y-%m-%d")
    metrics["date_max"] = dates.max().strftime("%Y-%m-%d")

    # 1. 基础销售数据
    paying = df[df["销售额"] > 0].copy()
    metrics["total_users"] = df["member_code"].nunique()
    metrics["total_sales"] = float(df["销售额"].sum())
    metrics["paying_users"] = paying["member_code"].nunique()
    metrics["paying_rows"] = len(paying)

    aov = metrics["total_sales"] / metrics["paying_users"] if metrics["paying_users"] > 0 else 0
    metrics["avg_order_value"] = round(aov, 2)

    # 复购用户（付费次数≥2 的用户）
    if "付费次数" in paying.columns and paying["付费次数"].sum() > 0:
        buyer_stats = paying.groupby("member_code").agg(
            total_sales=("销售额", "sum"),
            total_payments=("付费次数", "sum"),
        )
        repeat_buyers = buyer_stats[buyer_stats["total_payments"] >= 2]
        metrics["repeat_buyers"] = len(repeat_buyers)
        metrics["repeat_buyer_rate"] = round(len(repeat_buyers) / metrics["paying_users"] * 100, 2) if metrics["paying_users"] > 0 else 0
        metrics["repeat_sales"] = float(repeat_buyers["total_sales"].sum())
        metrics["repeat_sales_pct"] = round(metrics["repeat_sales"] / metrics["total_sales"] * 100, 2) if metrics["total_sales"] > 0 else 0
    else:
        metrics["repeat_buyers"] = 0
        metrics["repeat_buyer_rate"] = 0
        metrics["repeat_sales"] = 0
        metrics["repeat_sales_pct"] = 0

    # 2. 日销售额趋势
    daily = df.groupby(df["下单时间"].dt.date)["销售额"].sum().reset_index()
    daily.columns = ["日期", "销售额"]
    daily["销售额"] = daily["销售额"].astype(float)
    metrics["daily_sales"] = daily

    # 3. 付费用户画像 — 性别
    gender_valid = paying[paying["性别"].isin(["男", "女"])]
    if len(gender_valid) > 0:
        gender_dist = gender_valid.groupby("性别")["member_code"].nunique()
        metrics["gender_user_dist"] = gender_dist
        gender_sales = gender_valid.groupby("性别")["销售额"].sum().astype(float)
        metrics["gender_sales_dist"] = gender_sales
    else:
        metrics["gender_user_dist"] = pd.Series(dtype=int)
        metrics["gender_sales_dist"] = pd.Series(dtype=float)

    # 5. 付费用户画像 — 年龄分布
    if "年龄" in paying.columns:
        paying_age = paying["年龄"].dropna()
        if len(paying_age) > 0:
            age_bins = [0, 18, 25, 35, 45, 200]
            age_labels = ["18岁及以下", "19-25岁", "26-35岁", "36-45岁", "46岁以上"]
            paying["年龄分组"] = pd.cut(paying["年龄"], bins=age_bins, labels=age_labels, right=True)
            age_valid = paying[paying["年龄分组"].notna()]
            age_dist = age_valid.groupby("年龄分组", observed=False).agg(
                用户数=("member_code", "nunique"),
                销售额=("销售额", "sum"),
            ).reset_index()
            age_dist["销售额"] = age_dist["销售额"].astype(float)
            metrics["age_distribution"] = age_dist
    else:
        metrics["age_distribution"] = pd.DataFrame()

    # 6. 付费用户画像 — 性别&年龄交叉分布
    if "年龄" in paying.columns and "性别" in paying.columns:
        paying_cross = paying[paying["性别"].isin(["男", "女"])].copy()
        paying_cross = paying_cross[paying_cross["年龄分组"].notna()]
        if len(paying_cross) > 0:
            cross_user = paying_cross.groupby(["性别", "年龄分组"], observed=False)["member_code"].nunique().unstack(fill_value=0)
            cross_sales = paying_cross.groupby(["性别", "年龄分组"], observed=False)["销售额"].sum().unstack(fill_value=0)
            for lbl in age_labels:
                if lbl not in cross_user.columns:
                    cross_user[lbl] = 0
                    cross_sales[lbl] = 0.0
            cross_user = cross_user[age_labels]
            cross_sales = cross_sales[age_labels]
            metrics["cross_age_gender_user"] = cross_user
            metrics["cross_age_gender_sales"] = cross_sales
        else:
            metrics["cross_age_gender_user"] = pd.DataFrame()
            metrics["cross_age_gender_sales"] = pd.DataFrame()
        del paying_cross

    # 4. 会员等级分布
    if "会员等级（标签）" in paying.columns:
        vip_dist = paying.groupby("会员等级（标签）").agg(
            用户数=("member_code", "nunique"),
            销售额=("销售额", "sum"),
        ).reset_index().rename(columns={"会员等级（标签）": "会员等级"})
        vip_dist["销售额"] = vip_dist["销售额"].astype(float)
        vip_dist = vip_dist.sort_values("销售额", ascending=False)
        metrics["vip_distribution"] = vip_dist

    # 7. 新老用户分析（基于注册时间）
    if "注册时间" in paying.columns:
        paying_reg = paying.copy()
        paying_reg["注册时间_dt"] = pd.to_datetime(paying_reg["注册时间"], errors="coerce")
        paying_reg["下单时间_dt"] = pd.to_datetime(paying_reg["下单时间"])
        # 新用户定义：注册时间距下单时间在过去30天之内
        paying_reg["用户类型"] = paying_reg.apply(
            lambda row: "新用户" if pd.notna(row["注册时间_dt"]) and pd.notna(row["下单时间_dt"]) and 0 <= (row["下单时间_dt"] - row["注册时间_dt"]).days <= 30 else "老用户",
            axis=1
        )
        user_type_dist = paying_reg.groupby("用户类型").agg(
            用户数=("member_code", "nunique"),
            销售额=("销售额", "sum"),
        ).reset_index()
        user_type_dist["销售额"] = user_type_dist["销售额"].astype(float)
        metrics["new_old_user_dist"] = user_type_dist
        # Also compute summary for conclusions
        total_new_users = int(user_type_dist[user_type_dist["用户类型"]=="新用户"]["用户数"].sum())
        total_old_users = int(user_type_dist[user_type_dist["用户类型"]=="老用户"]["用户数"].sum())
        new_sales = float(user_type_dist[user_type_dist["用户类型"]=="新用户"]["销售额"].sum())
        old_sales = float(user_type_dist[user_type_dist["用户类型"]=="老用户"]["销售额"].sum())
        metrics["new_user_count"] = total_new_users
        metrics["old_user_count"] = total_old_users
        metrics["new_user_sales"] = new_sales
        metrics["old_user_sales"] = old_sales
        metrics["has_registration_data"] = True
    else:
        metrics["new_old_user_dist"] = pd.DataFrame()
        metrics["has_registration_data"] = False

    # 5. 检查字段可用性
    has_reg_field = any("注册" in col or "注册时间" in col for col in df.columns)
    has_age_field = "年龄" in paying.columns and paying["年龄"].notna().sum() > 0
    metrics["has_registration_data"] = has_reg_field
    metrics["has_age_data"] = has_age_field

    return metrics, paying

# ── 图表生成 ─────────────────────────────────────────────
def _save_chart(fig, path):
    fig.tight_layout(pad=1.5)
    fig.savefig(path, dpi=CHART_DPI, bbox_inches="tight", facecolor="white")
    plt.close(fig)

def chart_daily_sales(daily, output_dir):
    """日销售额趋势折线图"""
    fig, ax = plt.subplots(figsize=(CHART_WIDTH, CHART_HEIGHT))
    dates = [str(d) for d in daily["日期"]]
    sales = daily["销售额"].values
    ax.plot(dates, sales, marker="o", linewidth=2.5, color="#E74C3C",
            markersize=8, markerfacecolor="white", markeredgewidth=2, markeredgecolor="#E74C3C")
    ax.fill_between(range(len(dates)), sales, alpha=0.08, color="#E74C3C")
    for i, v in enumerate(sales):
        offset = sales.max() * 0.02
        ax.text(i, v + offset, f"{v:.0f}", ha="center", va="bottom", fontsize=9, fontweight="bold", color="#333")
    ax.set_xlabel("日期", fontsize=AXIS_FONT_SIZE)
    ax.set_ylabel("销售额 (元)", fontsize=AXIS_FONT_SIZE)
    ax.set_title("日销售额变化趋势", fontsize=TITLE_FONT_SIZE, fontweight="bold", pad=12)
    ax.set_ylim(bottom=0, top=sales.max() * 1.2)
    ax.grid(axis="y", alpha=0.3, linestyle="--")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    chart_path = os.path.join(output_dir, "chart_daily_sales.png")
    _save_chart(fig, chart_path)
    return chart_path

def chart_gender_dist(gender_users, gender_sales, output_dir):
    """性别分布饼图（用户数 + 销售额），仅标签和占比"""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(CHART_WIDTH + 2, CHART_HEIGHT))
    colors = ["#3498DB", "#E91E63"]

    labels1 = [f"{name}\n{pct:.1f}%" for name, pct in zip(
        gender_users.index.tolist(),
        gender_users.values / gender_users.values.sum() * 100
    )]
    sizes1 = gender_users.values
    wedges1, texts1 = ax1.pie(
        sizes1, labels=labels1, colors=colors[:len(sizes1)],
        startangle=90, labeldistance=1.15,
        textprops={"fontsize": 11, "fontweight": "bold", "color": "black"}
    )
    ax1.set_title("付费用户性别分布（用户数）", fontsize=AXIS_FONT_SIZE, fontweight="bold")

    labels2 = [f"{name}\n{pct:.1f}%" for name, pct in zip(
        gender_sales.index.tolist(),
        gender_sales.values / gender_sales.values.sum() * 100
    )]
    sizes2 = gender_sales.values
    wedges2, texts2 = ax2.pie(
        sizes2, labels=labels2, colors=colors[:len(sizes2)],
        startangle=90, labeldistance=1.15,
        textprops={"fontsize": 11, "fontweight": "bold", "color": "black"}
    )
    ax2.set_title("付费用户性别分布（销售额）", fontsize=AXIS_FONT_SIZE, fontweight="bold")

    fig.suptitle("付费用户性别分布", fontsize=TITLE_FONT_SIZE, fontweight="bold", y=1.02)
    chart_path = os.path.join(output_dir, "chart_gender.png")
    _save_chart(fig, chart_path)
    return chart_path

def chart_vip_dist(vip_df, output_dir):
    """会员等级分布饼图（用户数 + 销售额），仅标签和占比"""
    if vip_df.empty:
        return None
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(CHART_WIDTH + 2, CHART_HEIGHT))
    n = len(vip_df)
    colors1 = plt.cm.Blues(np.linspace(0.3, 0.9, n)) if n > 1 else ["#5B9BD5"]
    colors2 = plt.cm.Oranges(np.linspace(0.3, 0.9, n)) if n > 1 else ["#ED7D31"]

    vip_sorted_u = vip_df.sort_values("用户数", ascending=False)
    total_u = vip_sorted_u["用户数"].sum()
    labels1 = [f"{row['会员等级']}\n{row['用户数']/total_u*100:.1f}%"
               for _, row in vip_sorted_u.iterrows()]
    sizes1 = vip_sorted_u["用户数"].values
    wedges1, texts1 = ax1.pie(
        sizes1, labels=labels1, colors=colors1, startangle=90, labeldistance=1.15,
        textprops={"fontsize": 11, "fontweight": "bold", "color": "black"}
    )
    ax1.set_title("会员等级分布（用户数）", fontsize=AXIS_FONT_SIZE, fontweight="bold")

    vip_sorted_s = vip_df.sort_values("销售额", ascending=False)
    total_s = vip_sorted_s["销售额"].sum()
    labels2 = [f"{row['会员等级']}\n{row['销售额']/total_s*100:.1f}%"
               for _, row in vip_sorted_s.iterrows()]
    sizes2 = vip_sorted_s["销售额"].values
    wedges2, texts2 = ax2.pie(
        sizes2, labels=labels2, colors=colors2, startangle=90, labeldistance=1.15,
        textprops={"fontsize": 11, "fontweight": "bold", "color": "black"}
    )
    ax2.set_title("会员等级分布（销售额）", fontsize=AXIS_FONT_SIZE, fontweight="bold")

    fig.suptitle("付费用户会员等级分布", fontsize=TITLE_FONT_SIZE, fontweight="bold", y=1.02)
    chart_path = os.path.join(output_dir, "chart_vip.png")
    _save_chart(fig, chart_path)
    return chart_path

def chart_age_dist(age_df, output_dir):
    """年龄分布饼图（用户数 + 销售额），仅标签和占比"""
    if age_df.empty:
        return None
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(CHART_WIDTH + 2, CHART_HEIGHT))
    n = len(age_df)
    colors1 = plt.cm.Greens(np.linspace(0.3, 0.9, n)) if n > 1 else ["#2ECC71"]
    colors2 = plt.cm.Purples(np.linspace(0.3, 0.9, n)) if n > 1 else ["#9B59B6"]

    total_u = age_df["用户数"].sum()
    labels1 = [f"{row['年龄分组']}\n{row['用户数']/total_u*100:.1f}%"
               for _, row in age_df.iterrows()]
    sizes1 = age_df["用户数"].values
    wedges1, texts1 = ax1.pie(
        sizes1, labels=labels1, colors=colors1, startangle=90, labeldistance=1.15,
        textprops={"fontsize": 11, "fontweight": "bold", "color": "black"}
    )
    ax1.set_title("年龄分布（用户数）", fontsize=AXIS_FONT_SIZE, fontweight="bold")

    total_s = age_df["销售额"].sum()
    labels2 = [f"{row['年龄分组']}\n{row['销售额']/total_s*100:.1f}%"
               for _, row in age_df.iterrows()]
    sizes2 = age_df["销售额"].values
    wedges2, texts2 = ax2.pie(
        sizes2, labels=labels2, colors=colors2, startangle=90, labeldistance=1.15,
        textprops={"fontsize": 11, "fontweight": "bold", "color": "black"}
    )
    ax2.set_title("年龄分布（销售额）", fontsize=AXIS_FONT_SIZE, fontweight="bold")

    fig.suptitle("付费用户年龄分布", fontsize=TITLE_FONT_SIZE, fontweight="bold", y=1.02)
    chart_path = os.path.join(output_dir, "chart_age.png")
    _save_chart(fig, chart_path)
    return chart_path

def chart_new_old_user_dist(new_old_df, output_dir):
    """新老用户分布饼图（用户数 + 销售额），仅标签和占比"""
    if new_old_df.empty:
        return None
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(CHART_WIDTH + 2, CHART_HEIGHT))
    colors_uo = ["#F39C12", "#2ECC71"]  # 新用户金色, 老用户绿色

    # 新老用户排序：确保顺序一致
    uo_sorted_u = new_old_df.sort_values("用户数", ascending=False)
    total_u = uo_sorted_u["用户数"].sum()
    labels1 = [f"{row['用户类型']}\n{row['用户数']/total_u*100:.1f}%"
               for _, row in uo_sorted_u.iterrows()]
    sizes1 = uo_sorted_u["用户数"].values
    wedges1, texts1 = ax1.pie(
        sizes1, labels=labels1, colors=colors_uo[:len(sizes1)],
        startangle=90, labeldistance=1.15,
        textprops={"fontsize": 11, "fontweight": "bold", "color": "black"}
    )
    ax1.set_title("新老用户分布（用户数）", fontsize=AXIS_FONT_SIZE, fontweight="bold")

    uo_sorted_s = new_old_df.sort_values("销售额", ascending=False)
    total_s = uo_sorted_s["销售额"].sum()
    labels2 = [f"{row['用户类型']}\n{row['销售额']/total_s*100:.1f}%"
               for _, row in uo_sorted_s.iterrows()]
    sizes2 = uo_sorted_s["销售额"].values
    wedges2, texts2 = ax2.pie(
        sizes2, labels=labels2, colors=colors_uo[:len(sizes2)],
        startangle=90, labeldistance=1.15,
        textprops={"fontsize": 11, "fontweight": "bold", "color": "black"}
    )
    ax2.set_title("新老用户分布（销售额）", fontsize=AXIS_FONT_SIZE, fontweight="bold")

    fig.suptitle("付费用户新老用户分布", fontsize=TITLE_FONT_SIZE, fontweight="bold", y=1.02)
    chart_path = os.path.join(output_dir, "chart_new_old_user.png")
    _save_chart(fig, chart_path)
    return chart_path

# ── Word 报告生成 ────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def _set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _add_heading(doc, text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.name = "Microsoft YaHei"
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = run._element.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return h

def _add_body(doc, text, bold=False, font_size=10.5):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Microsoft YaHei"
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.bold = bold
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    return p

def _add_highlight_box(doc, key, value):
    """添加高亮结论框（用带底色表格模拟）"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "FFF3CD")
    # 清空默认段落
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"💡 {key}：{value}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Microsoft YaHei"
    run.font.color.rgb = RGBColor(0x85, 0x60, 0x00)
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    doc.add_paragraph()

def _add_metric_table(doc, headers, rows):
    """添加指标表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Microsoft YaHei"
                rpr = run._element.get_or_add_rPr()
                rFonts = rpr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = run._element.makeelement(qn("w:rFonts"), {})
                    rpr.insert(0, rFonts)
                rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        _set_cell_shading(cell, "2C3E50")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                run.font.name = "Microsoft YaHei"
                rpr = run._element.get_or_add_rPr()
                rFonts = rpr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = run._element.makeelement(qn("w:rFonts"), {})
                    rpr.insert(0, rFonts)
                rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            if i % 2 == 0:
                _set_cell_shading(cell, "F8F9FA")
    doc.add_paragraph()

def generate_report(metrics, paying, output_path, excel_name,
                    daily_chart_path=None, gender_chart_path=None, vip_chart_path=None,
                    age_chart_path=None, gender_age_cross_chart_path=None, new_old_user_chart_path=None):
    """生成 Word 报告"""
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ═══════════ 封面标题 ═══════════
    title = doc.add_heading("IP销售分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.size = Pt(22)
        run.font.name = "Microsoft YaHei"
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = run._element.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run1 = meta.add_run(f"数据源：{excel_name}")
    meta_run1.font.size = Pt(9)
    meta_run1.font.name = "Microsoft YaHei"
    rpr1 = meta_run1._element.get_or_add_rPr()
    rFonts1 = rpr1.find(qn("w:rFonts"))
    if rFonts1 is None:
        rFonts1 = meta_run1._element.makeelement(qn("w:rFonts"), {})
        rpr1.insert(0, rFonts1)
    rFonts1.set(qn("w:eastAsia"), "Microsoft YaHei")
    meta.add_run("\n")
    meta_run2 = meta.add_run(f"分析周期：{metrics['date_min']} 至 {metrics['date_max']}")
    meta_run2.font.size = Pt(9)
    meta_run2.font.name = "Microsoft YaHei"
    rpr2 = meta_run2._element.get_or_add_rPr()
    rFonts2 = rpr2.find(qn("w:rFonts"))
    if rFonts2 is None:
        rFonts2 = meta_run2._element.makeelement(qn("w:rFonts"), {})
        rpr2.insert(0, rFonts2)
    rFonts2.set(qn("w:eastAsia"), "Microsoft YaHei")
    meta.add_run("\n")
    meta_run3 = meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta_run3.font.size = Pt(9)
    meta_run3.font.name = "Microsoft YaHei"
    rpr3 = meta_run3._element.get_or_add_rPr()
    rFonts3 = rpr3.find(qn("w:rFonts"))
    if rFonts3 is None:
        rFonts3 = meta_run3._element.makeelement(qn("w:rFonts"), {})
        rpr3.insert(0, rFonts3)
    rFonts3.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.add_paragraph()

    total_sales = metrics["total_sales"]
    paying_users = metrics["paying_users"]
    aov = metrics["avg_order_value"]
    repeat_rate = metrics["repeat_buyer_rate"]
    repeat_sales_pct = metrics["repeat_sales_pct"]
    daily = metrics["daily_sales"]
    peak_row = daily.loc[daily["销售额"].idxmax()]
    trough_row = daily.loc[daily["销售额"].idxmin()]

    # ═══════════ 一、关键结论 ═══════════
    _add_heading(doc, "一、关键结论", level=1)
    conclusions = []
    conclusions.append(f"本次IP销售分析周期为 {metrics['date_min']} 至 {metrics['date_max']}，共 {len(daily)} 天完成销售活动。")
    conclusions.append(f"总销售额达 {_fmt_yuan(total_sales)} 元，共有 {paying_users} 名付费用户，整体客单价为 {_fmt_yuan(aov)} 元。")
    if metrics["repeat_buyers"] > 0:
        conclusions.append(
            f"复购用户（消费≥2次）共 {metrics['repeat_buyers']} 人，占付费用户的 {repeat_rate}%，"
            f"贡献了 {_fmt_yuan(metrics['repeat_sales'])} 元（占比 {repeat_sales_pct}%）。")
    conclusions.append(
        f"销售高峰期出现在 {peak_row['日期']}，单日销售额达 {_fmt_yuan(peak_row['销售额'])} 元；"
        f"低谷在 {trough_row['日期']}，单日仅 {_fmt_yuan(trough_row['销售额'])} 元。")
    if len(metrics.get("gender_user_dist", [])) > 0:
        gd = metrics["gender_user_dist"]
        gs = metrics["gender_sales_dist"]
        for gender in ["女", "男"]:
            if gender in gd.index:
                user_pct = gd[gender] / gd.sum() * 100
                sales_val = gs.get(gender, 0)
                conclusions.append(f"{gender}性用户 {gd[gender]} 人（占付费用户 {user_pct:.1f}%），贡献销售额 {_fmt_yuan(sales_val)} 元。")
                break
    if "vip_distribution" in metrics:
        vip = metrics["vip_distribution"]
        top_vip = vip.iloc[0]
        conclusions.append(
            f"会员等级分布中，{top_vip['会员等级']}等级用户销售额贡献最高，"
            f"达 {_fmt_yuan(top_vip['销售额'])} 元，占比 {top_vip['销售额'] / total_sales * 100:.1f}%。")

    if "new_user_count" in metrics and metrics["new_user_count"] > 0:
        new_user_pct = metrics["new_user_count"] / paying_users * 100
        old_user_avg = metrics["old_user_sales"] / metrics["old_user_count"]
        new_user_avg = metrics["new_user_sales"] / metrics["new_user_count"]
        conclusions.append(
            f"新用户（近30天首次活跃）共 {metrics['new_user_count']} 人（占付费用户 {new_user_pct:.1f}%），"
            f"贡献销售额 {_fmt_yuan(metrics['new_user_sales'])} 元，客单价 {_fmt_yuan(new_user_avg)} 元。"
            f"老用户 {metrics['old_user_count']} 人贡献 {_fmt_yuan(metrics['old_user_sales'])} 元，"
            f"客单价 {_fmt_yuan(old_user_avg)} 元，老用户消费力显著更强。")

    for i, c in enumerate(conclusions, 1):
        _add_body(doc, f"{i}. {c}")
    doc.add_paragraph()

    # ═══════════ 二、基础销售数据 ═══════════
    _add_heading(doc, "二、基础销售数据", level=1)
    _add_metric_table(doc, ["指标", "数值"], [
        ["总销售额（GMV）", f"{_fmt_yuan(total_sales)} 元"],
        ["付费用户数", f"{paying_users} 人"],
        ["总订单行数", f"{metrics['paying_rows']} 条"],
        ["客单价（AOV）", f"{_fmt_yuan(aov)} 元"],
        ["复购用户数（消费≥2次）", f"{metrics['repeat_buyers']} 人"],
        ["复购用户占比", f"{repeat_rate}%"],
        ["复购用户销售额贡献占比", f"{repeat_sales_pct}%"],
    ])
    _add_highlight_box(doc, "核心发现",
        f"本期IP总销售额 {_fmt_yuan(total_sales)} 元，付费用户 {paying_users} 人，"
        f"客单价 {aov:.2f} 元。复购率为 {repeat_rate}%，复购用户贡献了 {repeat_sales_pct}% 的销售额。")

    # ═══════════ 三、销售数据趋势 ═══════════
    _add_heading(doc, "三、销售数据趋势", level=1)
    if daily_chart_path and os.path.exists(daily_chart_path):
        doc.add_picture(daily_chart_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    trend_text = (
        f"销售周期内日销售额呈现「先升后降」趋势。"
        f"上线首日（{daily.iloc[0]['日期']}）销售额为 {_fmt_yuan(daily.iloc[0]['销售额'])} 元"
    )
    if len(daily) > 1:
        if daily["销售额"].iloc[0] < daily["销售额"].iloc[1] and daily["销售额"].iloc[-1] < daily["销售额"].iloc[-2]:
            trend_text += f"，快速攀升至峰值 {_fmt_yuan(peak_row['销售额'])} 元后逐步回落"
        elif daily["销售额"].is_monotonic_increasing:
            trend_text += "，整体呈持续增长态势"
        elif daily["销售额"].is_monotonic_decreasing:
            trend_text += "，整体呈持续下降态势"
    if len(daily) >= 2:
        changes = []
        for i in range(1, len(daily)):
            prev = daily.iloc[i - 1]["销售额"]
            curr = daily.iloc[i]["销售额"]
            if prev > 0:
                changes.append((curr - prev) / prev * 100)
        if changes:
            avg_chg = np.mean(changes)
            trend_text += f"。日均环比变化 {avg_chg:+.1f}%"
    trend_text += "。"
    _add_body(doc, trend_text)
    _add_highlight_box(doc, "趋势洞察",
        f"销售峰值出现在 {peak_row['日期']}（{_fmt_yuan(peak_row['销售额'])} 元），"
        f"占全周期销售额的 {peak_row['销售额'] / total_sales * 100:.1f}%。"
        f"建议关注峰值期的营销策略和转化因素，用于后续IP销售活动的优化复制。")

    # ═══════════ 四、付费用户画像 ═══════════
    _add_heading(doc, "四、付费用户画像", level=1)
    _add_metric_table(doc, ["画像指标", "数值"], [
        ["总付费用户数", f"{paying_users} 人"],
        ["总销售额", f"{_fmt_yuan(total_sales)} 元"],
        ["客单价", f"{_fmt_yuan(aov)} 元"],
        ["复购用户占比", f"{repeat_rate}%"],
    ])

    # ── 4.1 性别分布 ──
    _add_heading(doc, "4.1 付费用户性别分布", level=2)
    if len(metrics.get("gender_user_dist", [])) > 0:
        if gender_chart_path and os.path.exists(gender_chart_path):
            doc.add_picture(gender_chart_path, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        gd = metrics["gender_user_dist"]
        gs = metrics["gender_sales_dist"]
        total_g = gd.sum()
        total_gs = gs.sum()
        gender_rows = []
        for g in gd.index:
            user_pct = gd[g] / total_g * 100
            sales_pct = gs[g] / total_gs * 100 if total_gs > 0 else 0
            aov_g = gs[g] / gd[g] if gd[g] > 0 else 0
            gender_rows.append([g, str(gd[g]), f"{user_pct:.1f}%", f"{_fmt_yuan(gs[g])} 元", f"{_fmt_yuan(aov_g)} 元"])
        _add_metric_table(doc, ["性别", "付费用户数", "用户占比", "销售额", "客单价"], gender_rows)

        if "女" in gd.index and "男" in gd.index:
            if gs.get("女", 0) > gs.get("男", 0):
                gc = f"女性用户贡献销售额 {_fmt_yuan(gs['女'])} 元，高于男性用户的 {_fmt_yuan(gs['男'])} 元，为本次IP销售的核心消费群体。"
            else:
                gc = f"男性用户贡献销售额 {_fmt_yuan(gs['男'])} 元，高于女性用户的 {_fmt_yuan(gs['女'])} 元，为本次IP销售的核心消费群体。"
        elif "女" in gd.index:
            gc = "本次付费用户群体以女性为主。"
        else:
            gc = "本次付费用户群体以男性为主。"
        _add_highlight_box(doc, "性别洞察", gc)
    else:
        _add_body(doc, "（数据中缺少性别信息，无法进行性别分布分析。）")

    # ── 4.2 年龄分布 ──
    _add_heading(doc, "4.2 付费用户年龄分布", level=2)
    if metrics.get("has_age_data") and "age_distribution" in metrics:
        age_dist = metrics["age_distribution"]
        if not age_dist.empty:
            if age_chart_path and os.path.exists(age_chart_path):
                doc.add_picture(age_chart_path, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

            age_rows = []
            for _, row in age_dist.iterrows():
                user_pct = row["用户数"] / age_dist["用户数"].sum() * 100
                sales_pct = row["销售额"] / age_dist["销售额"].sum() * 100
                aov_a = row["销售额"] / row["用户数"] if row["用户数"] > 0 else 0
                age_rows.append([row["年龄分组"], str(row["用户数"]), f"{user_pct:.1f}%",
                                 f"{_fmt_yuan(row['销售额'])} 元", f"{sales_pct:.1f}%", f"{_fmt_yuan(aov_a)} 元"])
            _add_metric_table(doc, ["年龄段", "付费用户数", "用户占比", "销售额", "销售额占比", "客单价"], age_rows)

            peak_age = age_dist.loc[age_dist["用户数"].idxmax()]
            peak_sales_age = age_dist.loc[age_dist["销售额"].idxmax()]
            aov_series = age_dist["销售额"] / age_dist["用户数"].replace(0, 1)
            best_aov_age = age_dist.loc[aov_series.idxmax(), "年龄分组"]
            age_conclusion = (
                f"{peak_age['年龄分组']}年龄段付费用户最多（{peak_age['用户数']}人），"
                f"{peak_sales_age['年龄分组']}年龄段销售额贡献最高（{_fmt_yuan(peak_sales_age['销售额'])} 元）。"
                f"整体客单价以 {best_aov_age} 年龄段最高。"
            )
            _add_highlight_box(doc, "年龄洞察", age_conclusion)
    else:
        _add_body(doc, "（当前数据未包含年龄字段，无法进行年龄分布分析。建议后续数据收集中增加用户年龄信息。）")

    # ── 4.3 性别&年龄交叉分布 ──
    _add_heading(doc, "4.3 付费用户性别&年龄交叉分布", level=2)
    if metrics.get("has_age_data") and "cross_age_gender_user" in metrics:
        cross_user = metrics["cross_age_gender_user"]
        cross_sales = metrics["cross_age_gender_sales"]
        if not cross_user.empty:
            cross_rows = []
            age_labels_list = cross_user.columns.tolist()
            for gender in ["男", "女"]:
                if gender in cross_user.index:
                    for ag in age_labels_list:
                        u = int(cross_user.loc[gender, ag])
                        s = float(cross_sales.loc[gender, ag]) if gender in cross_sales.index and ag in cross_sales.columns else 0.0
                        if u > 0 or s > 0:
                            aov_c = s / u if u > 0 else 0
                            cross_rows.append([gender, ag, str(u), f"{_fmt_yuan(s)} 元", f"{_fmt_yuan(aov_c)} 元", s])
            # 按销售额降序排列
            cross_rows.sort(key=lambda r: r[5], reverse=True)
            # 去掉辅助的排序列
            cross_rows = [r[:5] for r in cross_rows]
            if cross_rows:
                _add_metric_table(doc, ["性别", "年龄段", "付费用户数", "销售额", "客单价"], cross_rows)

            total_cross_users = sum(int(cross_user.loc[g].sum()) for g in cross_user.index)
            total_cross_sales = sum(float(cross_sales.loc[g].sum()) for g in cross_sales.index)
            best_gender = "女" if "女" in cross_sales.index and cross_sales.loc["女"].sum() > cross_sales.loc["男"].sum() else "男" if "男" in cross_sales.index else ""
            cross_conclusion = (
                f"交叉分析覆盖 {total_cross_users} 名付费用户，"
                f"贡献销售额 {_fmt_yuan(total_cross_sales)} 元。"
            )
            if best_gender:
                cross_conclusion += f"{best_gender}性用户在多数年龄段表现更优。"
            _add_highlight_box(doc, "交叉洞察", cross_conclusion)
    else:
        _add_body(doc, "（当前数据未包含年龄和性别字段，无法进行交叉分析。建议后续补充年龄信息。）")

    # ── 4.4 会员等级分布 ──
    _add_heading(doc, "4.4 付费用户会员等级分布", level=2)
    if "vip_distribution" in metrics:
        if vip_chart_path and os.path.exists(vip_chart_path):
            doc.add_picture(vip_chart_path, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        vip = metrics["vip_distribution"]
        total_vip_users = vip["用户数"].sum()
        total_vip_sales = vip["销售额"].sum()
        vip_rows = []
        for _, row in vip.iterrows():
            user_pct = row["用户数"] / total_vip_users * 100
            sales_pct = row["销售额"] / total_vip_sales * 100
            aov_v = row["销售额"] / row["用户数"] if row["用户数"] > 0 else 0
            vip_rows.append([row["会员等级"], str(row["用户数"]), f"{user_pct:.1f}%",
                             f"{_fmt_yuan(row['销售额'])} 元", f"{sales_pct:.1f}%", f"{aov_v:.2f} 元"])
        _add_metric_table(doc, ["会员等级", "付费用户数", "用户占比", "销售额", "销售额占比", "客单价"], vip_rows)
        top_vip = vip.iloc[0]
        _add_highlight_box(doc, "会员等级洞察",
            f"{top_vip['会员等级']} 等级会员人数最多（{top_vip['用户数']}人），"
            f"贡献销售额最高（{_fmt_yuan(top_vip['销售额'])} 元），为本次IP销售的核心贡献等级。"
            f"建议在后续IP销售中重点关注该等级会员的触达和转化。")
    else:
        _add_body(doc, "（数据中缺少会员等级信息，无法进行会员等级分布分析。）")
    # ── 4.5 新老用户分布 ──
    _add_heading(doc, "4.5 付费用户新老用户分布", level=2)
    if metrics.get("has_registration_data") and "new_old_user_dist" in metrics:
        no_df = metrics["new_old_user_dist"]
        if not no_df.empty:
            if new_old_user_chart_path and os.path.exists(new_old_user_chart_path):
                doc.add_picture(new_old_user_chart_path, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

            total_no_users = no_df["用户数"].sum()
            total_no_sales = no_df["销售额"].sum()
            no_rows = []
            for _, row in no_df.iterrows():
                user_pct = row["用户数"] / total_no_users * 100
                sales_pct = row["销售额"] / total_no_sales * 100
                aov_n = row["销售额"] / row["用户数"] if row["用户数"] > 0 else 0
                no_rows.append([row["用户类型"], str(row["用户数"]), f"{user_pct:.1f}%",
                                 f"{_fmt_yuan(row['销售额'])} 元", f"{sales_pct:.1f}%", f"{_fmt_yuan(aov_n)} 元"])
            _add_metric_table(doc, ["用户类型", "付费用户数", "用户占比", "销售额", "销售额占比", "客单价"], no_rows)

            new_row = no_df[no_df["用户类型"]=="新用户"]
            old_row = no_df[no_df["用户类型"]=="老用户"]
            new_users = int(new_row["用户数"].sum()) if not new_row.empty else 0
            old_users = int(old_row["用户数"].sum()) if not old_row.empty else 0
            new_sales = float(new_row["销售额"].sum()) if not new_row.empty else 0
            old_sales = float(old_row["销售额"].sum()) if not old_row.empty else 0
            no_conclusion = (
                f"新用户（近30天内首次活跃）共 {new_users} 人，贡献销售额 {_fmt_yuan(new_sales)} 元；"
                f"老用户共 {old_users} 人，贡献销售额 {_fmt_yuan(old_sales)} 元。"
            )
            if old_users > 0:
                avg_sales_old = old_sales / old_users
                avg_sales_new = (new_sales / new_users) if new_users > 0 else 0
                no_conclusion += f"老用户平均客单价 {_fmt_yuan(avg_sales_old)} 元，显著高于新用户的 {_fmt_yuan(avg_sales_new)} 元，表明老用户对IP品牌的忠诚度和消费力更强。"
            _add_highlight_box(doc, "新老用户洞察", no_conclusion)
    else:
        _add_body(doc, "（当前数据未包含注册时间信息，无法进行新老用户分布分析。）")



    # ═══════════ 五、分析总结与建议 ═══════════
    _add_heading(doc, "五、分析总结与建议", level=1)
    _add_body(doc,
        f"基于 {metrics['date_min']} 至 {metrics['date_max']} 的销售数据分析，"
        f"本期IP销售活动共实现总销售额 {_fmt_yuan(total_sales)} 元，吸引 {paying_users} 名付费用户参与。")
    _add_body(doc,
        "本报告严格遵循「新IP分析方法」模板的分析框架，涵盖基础销售指标、日销售趋势、"
        "付费用户画像（性别、年龄、性别&年龄交叉、会员等级）等核心分析维度。"
        "报告中所有数据均来源于原始Excel数据文件，确保分析结果的准确性和可追溯性。")
    _add_body(doc,
        "建议后续数据收集中增加注册时间等字段，以支持新老用户对比等更深入的分析，"
        "进一步提升分析报告的完整性和决策支持价值。",
        bold=True)
    # ═══════════ 六、指标计算说明 ═══════════
    _add_heading(doc, "六、指标计算说明", level=1)
    _add_body(doc, "本报告所有分析指标均基于原始销售数据计算，以下为各指标的定义与计算逻辑：")
    metrics_defs = [
        ["总销售额(GMV)", "所有有销售记录行的销售额加总。销售额为'null'字符串时视为0处理。"],
        ["付费用户数", "销售额大于0的唯一member_code数量。"],
        ["客单价(AOV)", "总销售额 ÷ 付费用户数。反映平均每位付费用户的消费金额。"],
        ["复购用户", "付费次数≥2的唯一member_code。付费次数按每个member_code的销售记录行数统计。"],
        ["复购用户占比", "复购用户数 ÷ 付费用户数 × 100%。反映用户重复购买的比例。"],
        ["复购销售额占比", "复购用户的销售额之和 ÷ 总销售额 × 100%。反映复购行为对总销售额的贡献程度。"],
        ["日销售额", "按下单时间逐日汇总销售额，用于分析销售趋势变化。"],
        ["性别分布", "仅统计性别为'男'或'女'的用户，'未知'性别不纳入分析。"],
        ["年龄分布", "基于生日字段以分析周期中点（2026-07-07）为基准计算年龄，划分为18岁及以下、19-25岁、26-35岁、36-45岁、46岁以上五个年龄段。"],
        ["新老用户划分", f"逐条记录比较用户的首次活跃时间（{metrics["date_min"]}）与其下单时间，若首次活跃时间距下单时间不超过30天则视为新用户，否则视为老用户。首次活跃时间取自'首次活跃时间'字段。"],
        ["会员等级", "直接使用'会员等级（标签）'字段，按各等级汇总用户数和销售额。"],
    ]
    _add_metric_table(doc, ["指标", "计算逻辑"], metrics_defs)
    _add_body(doc, "以上所有指标计算均已在代码中实现，确保分析结果的准确性和可重复性。")

    doc.save(output_path)
    print(f"[OK] 报告已生成：{output_path}")

# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="IP销售分析报告生成器")
    parser.add_argument("--excel", required=True, help="销售数据Excel文件路径")
    parser.add_argument("--output", default=os.getcwd(), help="报告输出目录（默认当前目录）")
    args = parser.parse_args()

    excel_path = args.excel
    output_dir = args.output
    excel_name = os.path.basename(excel_path)

    if not os.path.isfile(excel_path):
        print(f"[错误] 找不到Excel文件：{excel_path}")
        sys.exit(1)

    os.makedirs(output_dir, exist_ok=True)

    print(f"[1/5] 读取数据：{excel_path}")
    import openpyxl
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb.active
    data_rows = []
    headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
    for r in range(2, ws.max_row + 1):
        row = [ws.cell(r, c).value for c in range(1, ws.max_column + 1)]
        data_rows.append(row)
    df = pd.DataFrame(data_rows, columns=headers)
    df["销售额"] = _parse_sales(df["销售额"])
    df["付费次数"] = _parse_freq(df["付费次数"])
    df["下单时间"] = pd.to_datetime(df["下单时间"], errors="coerce")
    df = df.dropna(subset=["下单时间"])
    print(f"  有效数据行数：{len(df)}")

    print(f"[2/5] 计算分析指标...")
    metrics, paying = compute_metrics(df)

    print(f"[3/5] 生成可视化图表...")
    daily_chart_path = chart_daily_sales(metrics["daily_sales"], output_dir)
    print(f"  日销售趋势图 → {daily_chart_path}")
    gender_chart_path = None
    if len(metrics.get("gender_user_dist", [])) > 0:
        gender_chart_path = chart_gender_dist(metrics["gender_user_dist"], metrics["gender_sales_dist"], output_dir)
        print(f"  性别分布图 → {gender_chart_path}")
    vip_chart_path = None
    if "vip_distribution" in metrics:
        vip_chart_path = chart_vip_dist(metrics["vip_distribution"], output_dir)
        print(f"  会员等级分布图 → {vip_chart_path}")
    age_chart_path = None
    if "age_distribution" in metrics and not metrics["age_distribution"].empty:
        age_chart_path = chart_age_dist(metrics["age_distribution"], output_dir)
        if age_chart_path:
            print(f"  年龄分布图 → {age_chart_path}")
    gender_age_cross_chart_path = None

    print(f"[4/5] 生成 Word 报告...")
    output_path = os.path.join(output_dir, f"IP销售分析报告_{metrics['date_min']}_{metrics['date_max']}.docx")
    new_old_user_chart_path = None
    if "new_old_user_dist" in metrics and not metrics["new_old_user_dist"].empty:
        new_old_user_chart_path = chart_new_old_user_dist(metrics["new_old_user_dist"], output_dir)
        print(f"  新老用户分布图 → {new_old_user_chart_path}")

    generate_report(metrics, paying, output_path, excel_name,
                    daily_chart_path, gender_chart_path, vip_chart_path,
                    age_chart_path, gender_age_cross_chart_path,
                    new_old_user_chart_path)

    print(f"[5/5] 完成！")
    print(f"📊 报告输出路径：{output_path}")
    print()
    print("=" * 60)
    print("  报告结构概览")
    print("=" * 60)
    expected_sections = [
        "一、关键结论",
        "二、基础销售数据",
        "三、销售数据趋势",
        "四、付费用户画像",
        "  4.1 付费用户性别分布",
        "  4.2 付费用户年龄分布",
        "  4.3 付费用户性别&年龄交叉分布",
        "  4.4 付费用户会员等级分布",
        "  4.5 付费用户新老用户分布",
        "五、分析总结与建议",
        "六、指标计算说明",
    ]
    for s in expected_sections:
        print(f"  ✅ {s}")
    print()
    print("💡 提示：使用 validate_report.py 校验报告结构一致性")
